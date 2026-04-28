"""DOCX parsing and in-place editing utilities for resume tailoring."""

from __future__ import annotations

import difflib
import logging
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document as load_document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph

from openai_client import TailoredResume

LOGGER = logging.getLogger(__name__)

SUMMARY_HEADINGS = {
    "summary",
    "professionalsummary",
    "profile",
}

EXPERIENCE_HEADINGS = {
    "experience",
    "professionalexperience",
    "workexperience",
    "employmenthistory",
}

OTHER_SECTION_HEADINGS = {
    "education",
    "skills",
    "technicalskills",
    "projects",
    "certifications",
    "awards",
    "volunteer",
    "volunteerexperience",
    "publications",
    "additionalinformation",
}


@dataclass(slots=True)
class ExperienceSection:
    """Parsed experience block in the DOCX resume."""

    company: str
    role: str
    bullets: list[str]
    company_index: int
    role_indices: list[int]
    bullet_indices: list[int]


@dataclass(slots=True)
class ResumeStructure:
    """Parsed resume structure with paragraph anchors for mutation."""

    path: Path
    document: DocxDocument
    summary_index: int
    summary_text: str
    experiences: list[ExperienceSection]

    @property
    def paragraphs(self) -> list[Paragraph]:
        return self.document.paragraphs

    def full_text(self) -> str:
        return "\n".join(paragraph.text.strip() for paragraph in self.paragraphs if paragraph.text.strip())

    def editable_payload(self) -> dict[str, object]:
        return {
            "summary": self.summary_text,
            "experiences": [
                {
                    "company": experience.company,
                    "role": experience.role,
                    "bullets": experience.bullets,
                }
                for experience in self.experiences
            ],
        }


class DocxEditingError(RuntimeError):
    """Raised when the resume template cannot be parsed or edited safely."""


def load_resume_structure(path: str | Path) -> ResumeStructure:
    """Parse the resume DOCX into editable sections."""

    path = Path(path).expanduser().resolve()
    document = load_document(path)
    paragraphs = document.paragraphs

    experience_heading_index = _find_heading_index(paragraphs, EXPERIENCE_HEADINGS)
    if experience_heading_index is None:
        raise DocxEditingError(
            f"Could not find an experience heading. Expected one of: {sorted(EXPERIENCE_HEADINGS)}."
        )

    summary_index = _find_summary_index(paragraphs, experience_heading_index)
    if summary_index is None:
        raise DocxEditingError(
            "Could not identify the summary paragraph. Add a Summary heading or place the summary directly before the experience section."
        )

    experience_end_index = _find_next_section_index(paragraphs, experience_heading_index + 1)
    experiences = _parse_experiences(
        paragraphs=paragraphs,
        start_index=experience_heading_index + 1,
        end_index=experience_end_index,
    )
    if not experiences:
        raise DocxEditingError("Could not identify any experience blocks under the experience heading.")

    LOGGER.info(
        "Parsed resume with %s experience section(s) from %s",
        len(experiences),
        path,
    )
    return ResumeStructure(
        path=path,
        document=document,
        summary_index=summary_index,
        summary_text=paragraphs[summary_index].text.strip(),
        experiences=experiences,
    )


def validate_tailoring(structure: ResumeStructure, tailoring: TailoredResume) -> None:
    """Ensure the model output can be applied without damaging the template."""

    if len(tailoring.experiences) != len(structure.experiences):
        raise DocxEditingError(
            "Model output changed the number of experience sections, which is not allowed."
        )

    for index, (source, generated) in enumerate(zip(structure.experiences, tailoring.experiences, strict=True)):
        if _normalize_text(source.company) != _normalize_text(generated.company):
            raise DocxEditingError(
                f"Experience {index} company mismatch: expected '{source.company}', got '{generated.company}'."
            )

        if not _roles_match(source.role, generated.role):
            raise DocxEditingError(
                f"Experience {index} role mismatch: expected '{source.role}', got '{generated.role}'."
            )

        if len(generated.bullets) != len(source.bullets):
            raise DocxEditingError(
                f"Experience {index} bullet count changed from {len(source.bullets)} to {len(generated.bullets)}."
            )

        expected_order = set(range(len(source.bullets)))
        generated_order = set(generated.bullet_order)
        if generated_order != expected_order or len(generated.bullet_order) != len(source.bullets):
            raise DocxEditingError(
                f"Experience {index} bullet_order must be a permutation of 0..{len(source.bullets) - 1}."
            )


def apply_tailoring(structure: ResumeStructure, tailoring: TailoredResume) -> None:
    """Mutate the loaded document in place."""

    validate_tailoring(structure, tailoring)

    replace_paragraph_text(structure.paragraphs[structure.summary_index], tailoring.summary.strip())

    for source, generated in zip(structure.experiences, tailoring.experiences, strict=True):
        reordered_bullets = [generated.bullets[index].strip() for index in generated.bullet_order]
        for paragraph_index, bullet_text in zip(source.bullet_indices, reordered_bullets, strict=True):
            replace_paragraph_text(structure.paragraphs[paragraph_index], bullet_text)

    structure.summary_text = tailoring.summary.strip()
    for source, generated in zip(structure.experiences, tailoring.experiences, strict=True):
        source.bullets = [generated.bullets[index].strip() for index in generated.bullet_order]


def save_resume(structure: ResumeStructure, output_path: str | Path) -> Path:
    """Write the edited document to disk."""

    output_path = Path(output_path).expanduser().resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    structure.document.save(output_path)
    LOGGER.info("Saved tailored DOCX to %s", output_path)
    return output_path


def build_diff_preview(structure: ResumeStructure, tailoring: TailoredResume) -> str:
    """Generate a unified diff for the planned content changes."""

    validate_tailoring(structure, tailoring)

    lines: list[str] = []
    lines.extend(_diff_block("Summary", structure.summary_text, tailoring.summary.strip()))

    for source, generated in zip(structure.experiences, tailoring.experiences, strict=True):
        before = "\n".join(source.bullets)
        after = "\n".join(generated.bullets[index].strip() for index in generated.bullet_order)
        heading = f"{source.company} | {source.role}"
        lines.extend(_diff_block(heading, before, after))

    return "\n".join(lines)


def replace_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
    """Replace paragraph content while preserving paragraph-level formatting."""

    new_text = new_text.strip()
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
        return

    paragraph.add_run(new_text)


def _parse_experiences(
    *,
    paragraphs: list[Paragraph],
    start_index: int,
    end_index: int | None,
) -> list[ExperienceSection]:
    """Parse experience blocks from the experience section."""

    stop_index = end_index if end_index is not None else len(paragraphs)
    nonempty_indices = [index for index in range(start_index, stop_index) if paragraphs[index].text.strip()]
    experiences: list[ExperienceSection] = []
    cursor = 0
    current_company: str | None = None

    while cursor < len(nonempty_indices):
        current_index = nonempty_indices[cursor]
        current_paragraph = paragraphs[current_index]

        if _is_content_bullet(current_paragraph):
            cursor += 1
            continue

        heading_indices = [current_index]
        cursor += 1
        while cursor < len(nonempty_indices):
            next_index = nonempty_indices[cursor]
            next_paragraph = paragraphs[next_index]
            if _is_content_bullet(next_paragraph):
                break
            if _normalize_text(next_paragraph.text) in OTHER_SECTION_HEADINGS:
                break
            heading_indices.append(next_index)
            cursor += 1

        bullet_indices: list[int] = []
        while cursor < len(nonempty_indices):
            next_index = nonempty_indices[cursor]
            next_paragraph = paragraphs[next_index]
            if not _is_content_bullet(next_paragraph):
                break
            bullet_indices.append(next_index)
            cursor += 1

        if not bullet_indices:
            continue

        first_heading_text = paragraphs[heading_indices[0]].text.strip()
        if _is_company_heading_text(first_heading_text):
            company_index = heading_indices[0]
            current_company = first_heading_text
            role_indices = heading_indices[1:]
        elif current_company is not None:
            company_index = -1
            role_indices = heading_indices
        else:
            company_index = heading_indices[0]
            current_company = first_heading_text
            role_indices = heading_indices[1:]

        company_text = current_company or paragraphs[heading_indices[0]].text.strip()
        role_text = " | ".join(paragraphs[index].text.strip() for index in role_indices)
        experiences.append(
            ExperienceSection(
                company=company_text,
                role=role_text,
                bullets=[paragraphs[index].text.strip() for index in bullet_indices],
                company_index=company_index,
                role_indices=role_indices,
                bullet_indices=bullet_indices,
            )
        )

    return experiences


def _find_heading_index(paragraphs: list[Paragraph], headings: set[str]) -> int | None:
    for index, paragraph in enumerate(paragraphs):
        if _normalize_text(paragraph.text) in headings:
            return index
    return None


def _find_next_section_index(paragraphs: list[Paragraph], start_index: int) -> int | None:
    for index in range(start_index, len(paragraphs)):
        if _normalize_text(paragraphs[index].text) in OTHER_SECTION_HEADINGS:
            return index
    return None


def _next_nonempty_index(paragraphs: list[Paragraph], start_index: int) -> int | None:
    for index in range(start_index, len(paragraphs)):
        if paragraphs[index].text.strip():
            return index
    return None


def _find_summary_index(paragraphs: list[Paragraph], experience_heading_index: int) -> int | None:
    summary_heading_index = _find_heading_index(paragraphs, SUMMARY_HEADINGS)
    if summary_heading_index is not None:
        return _next_nonempty_index(paragraphs, summary_heading_index + 1)

    for index in range(experience_heading_index - 1, -1, -1):
        paragraph = paragraphs[index]
        text = paragraph.text.strip()
        if not text:
            continue
        if _is_header_paragraph(paragraph):
            continue
        return index
    return None


def _is_bullet_paragraph(paragraph: Paragraph) -> bool:
    style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
    if "bullet" in style_name.lower() or "list" in style_name.lower():
        return True

    paragraph_properties = paragraph._p.pPr  # type: ignore[attr-defined]
    if paragraph_properties is not None and paragraph_properties.numPr is not None:
        return True

    text = paragraph.text.strip()
    return bool(re.match(r"^([-\u2022\u25cf\u25aa])\s+", text))


def _is_content_bullet(paragraph: Paragraph) -> bool:
    if not _is_bullet_paragraph(paragraph):
        return False
    return not _is_experience_header_candidate(paragraph)


def _is_experience_header_candidate(paragraph: Paragraph) -> bool:
    text = paragraph.text.strip()
    if not text:
        return False

    if "|" in text and len(text) <= 120:
        return True

    if len(text) <= 90 and not re.search(r"[.!?]$", text):
        return True

    return False


def _is_header_paragraph(paragraph: Paragraph) -> bool:
    text = paragraph.text.strip()
    if not text:
        return False

    if paragraph.alignment == 1:
        return True

    normalized = _normalize_text(text)
    if normalized in EXPERIENCE_HEADINGS or normalized in OTHER_SECTION_HEADINGS or normalized in SUMMARY_HEADINGS:
        return True

    if text.upper() == text and len(text) <= 40:
        return True

    return False


def _is_company_heading_text(text: str) -> bool:
    if "|" not in text:
        return False

    if _contains_date_range(text):
        return False

    return len(text) <= 120


def _contains_date_range(text: str) -> bool:
    month_pattern = (
        r"(jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)"
        r"[a-z]*\.?\s+\d{4}"
    )
    return bool(re.search(month_pattern, text, re.IGNORECASE))


def _normalize_text(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", value.lower())


def _roles_match(source_role: str, generated_role: str) -> bool:
    """Allow semantically equivalent role strings with optional extra subheading segments."""

    if _normalize_text(source_role) == _normalize_text(generated_role):
        return True

    source_parts = _role_parts(source_role)
    generated_parts = _role_parts(generated_role)
    if not source_parts or not generated_parts:
        return False

    common_length = min(len(source_parts), len(generated_parts))
    return source_parts[:common_length] == generated_parts[:common_length]


def _role_parts(value: str) -> list[str]:
    return [_normalize_text(part) for part in value.split("|") if _normalize_text(part)]


def _diff_block(label: str, before: str, after: str) -> list[str]:
    before_lines = before.splitlines() or [""]
    after_lines = after.splitlines() or [""]
    return list(
        difflib.unified_diff(
            before_lines,
            after_lines,
            fromfile=f"{label} (before)",
            tofile=f"{label} (after)",
            lineterm="",
        )
    )
